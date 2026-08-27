from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docs"
ASSETS = ROOT / "tmp" / "submission-assets"
OUT.mkdir(parents=True, exist_ok=True)
ASSETS.mkdir(parents=True, exist_ok=True)

NAVY = "143642"
TEAL = "0B6E69"
ORANGE = "E76F51"
INK = "243238"
MUTED = "66747A"
LIGHT = "EDF4F3"
PALE = "F6F8F8"
WHITE = "FFFFFF"
RULE = "D5E0DF"


def font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    ind = tbl_pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure(doc, label):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, TEAL),
        ("Heading 2", 13, 12, 6, TEAL),
        ("Heading 3", 12, 8, 4, NAVY),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"FURSHIELD  /  {label.upper()}")
    font(run, 8.5, True, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("FurShield  •  August 2026  •  ")
    font(run, 8.5, False, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def cover(doc, kicker, title, subtitle, metadata):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(kicker.upper())
    font(r, 10, True, ORANGE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(title)
    font(r, 29, True, NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run(subtitle)
    font(r, 14, False, TEAL)
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    table_geometry(table, [2700, 6660])
    set_repeat_table_header(table.rows[0])
    for row, (label, value) in zip(table.rows, metadata):
        shade(row.cells[0], LIGHT)
        for idx, text in enumerate((label, value)):
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            font(r, 10.5, idx == 0, NAVY if idx == 0 else INK)
    doc.add_paragraph()


def add_p(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_lead):])
        font(r)
    else:
        font(p.add_run(text))
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(item))


def steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        font(p.add_run(item))


def callout(doc, label, text, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    font(p.add_run(label + "  "), 10.5, True, TEAL)
    font(p.add_run(text), 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        shade(table.rows[0].cells[idx], TEAL)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(text), 9, True, WHITE)
    for values in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(values):
            if len(table.rows) % 2 == 1:
                shade(cells[idx], PALE)
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(text)), 8.6, False, INK)
        table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def diagram(path, title, nodes, arrows):
    img = Image.new("RGB", (1500, 780), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("arialbd.ttf", 42)
        body_font = ImageFont.truetype("arial.ttf", 25)
    except OSError:
        title_font = body_font = ImageFont.load_default()
    draw.text((60, 35), title, fill="#143642", font=title_font)
    for x, y, w, h, label, fill in nodes:
        draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=fill, outline="#0B6E69", width=4)
        box = draw.textbbox((0, 0), label, font=body_font)
        tx = x + (w - (box[2] - box[0])) / 2
        ty = y + (h - (box[3] - box[1])) / 2 - 3
        draw.multiline_text((tx, ty), label, fill="#143642", font=body_font, anchor="la", align="center")
    for x1, y1, x2, y2 in arrows:
        draw.line((x1, y1, x2, y2), fill="#E76F51", width=7)
        if x2 >= x1:
            draw.polygon([(x2, y2), (x2 - 18, y2 - 12), (x2 - 18, y2 + 12)], fill="#E76F51")
        else:
            draw.polygon([(x2, y2), (x2 + 18, y2 - 12), (x2 + 18, y2 + 12)], fill="#E76F51")
    img.save(path)


def add_figure(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(6.35))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    font(p.add_run(caption), 9, False, MUTED, italic=True)


def build_report():
    arch = ASSETS / "architecture.png"
    flow = ASSETS / "data-flow.png"
    diagram(arch, "Role-aware application architecture", [
        (70, 165, 260, 115, "Public\nvisitor", "#EDF4F3"),
        (70, 475, 260, 115, "Owner / Vet\n/ Shelter", "#FFF1EC"),
        (515, 315, 370, 135, "Next.js 16\nUI + server actions", "#EDF4F3"),
        (1090, 130, 315, 120, "MongoDB Atlas\ndomain data", "#EDF4F3"),
        (1090, 335, 315, 120, "Cloudinary\nimages + PDFs", "#FFF1EC"),
        (1090, 540, 315, 120, "Mailtrap SMTP\nemail capture", "#F6F8F8"),
    ], [(330, 222, 515, 355), (330, 532, 515, 410), (885, 365, 1090, 190), (885, 385, 1090, 395), (885, 410, 1090, 600)])
    diagram(flow, "Protected mutation and notification flow", [
        (45, 305, 260, 120, "Browser\nrequest", "#EDF4F3"),
        (390, 305, 260, 120, "Session + role\ncheck", "#FFF1EC"),
        (735, 305, 260, 120, "Validate +\nauthorize entity", "#EDF4F3"),
        (1080, 175, 330, 120, "Persist domain\nchange", "#EDF4F3"),
        (1080, 450, 330, 120, "Create in-app alert\n+ optional email", "#FFF1EC"),
    ], [(305, 365, 390, 365), (650, 365, 735, 365), (995, 350, 1080, 235), (995, 390, 1080, 510)])

    doc = Document()
    configure(doc, "Software Requirements Completion Report")
    cover(doc, "Full-stack application", "FurShield", "Software Requirements Completion Report", [
        ("Purpose", "Implementation traceability, architecture, data design, verification, and handoff"),
        ("Source", "PetCare Full-Stack App SRS (final)"),
        ("Stack", "Next.js 16, React 19, MongoDB Atlas, Cloudinary, signed cookie sessions"),
        ("Prepared", date.today().strftime("%d %B %Y")),
        ("Release status", "Local release candidate; deployment and MP4 are external release artifacts"),
    ])
    callout(doc, "Completion statement", "All mandatory application workflows in the SRS have an implemented local path. Optional family sharing and chatbot features remain intentionally excluded. Payment, delivery, and vet credential authentication are explicitly outside scope. Hosting and recording cannot be certified by a local build.")

    doc.add_heading("1. Problem and proposed solution", level=1)
    add_p(doc, "Pet-care information is often split across paper files, clinics, calendars, shelters, shops, and informal messages. That fragmentation makes reminders easier to miss and forces veterinarians and shelters to reconstruct context. FurShield centralizes these workflows around the animal while preserving role boundaries.")
    bullets(doc, [
        "Owners maintain multiple pets, rich health timelines, documents, appointments, adoption requests, cart activity, notifications, and reviews.",
        "Veterinarians publish availability, manage bookings, access only booked-pet histories, and record structured clinical observations and treatment.",
        "Shelters publish image-based profiles, maintain care logs, review adopter forms, reply, and finalize decisions.",
        "Visitors search care content, products, veterinarians, adoption listings, and public feedback without signing in.",
    ])

    doc.add_heading("2. Scope, assumptions, and constraints", level=1)
    matrix(doc, ["Class", "Decision", "Treatment"], [
        ("Mandatory application scope", "Owner, veterinarian, shelter, and public workflows", "Implemented in the local application"),
        ("Optional", "Family sharing; AI chatbot", "Not enabled; not required for acceptance"),
        ("Excluded by SRS", "Payments, physical delivery, vet credential authentication", "Deliberately absent"),
        ("External configuration", "Email provider, deployment target", "Environment-driven; in-app alerts work locally"),
        ("Submission artifact", "MP4 walkthrough", "Script supplied; recording requires an interactive release environment"),
    ], [1850, 4000, 3510])

    doc.add_heading("3. Architecture and system flow", level=1)
    add_figure(doc, arch, "Figure 1. Browser, application, data, media, and optional email boundaries.")
    add_p(doc, "The Next.js application uses server-rendered routes for data-aware screens, server actions for form mutations, and route handlers for signed uploads. Every protected route and mutation repeats session, role, and ownership checks; navigation visibility is not treated as authorization.")
    add_figure(doc, flow, "Figure 2. Authorization precedes persistence and user notification.")

    doc.add_heading("4. Database design", level=1)
    matrix(doc, ["Collection", "Key relationships", "Purpose"], [
        ("users", "role; owned pets/listings; reviews", "Identity, contact data, vet profile/availability, shelter profile"),
        ("pets", "ownerId", "Identity, gallery, allergies, ongoing conditions, microchip and weight"),
        ("healthrecords", "petId, ownerId, optional vetId", "Vaccines, allergies, illness, treatment, labs, documents, insurance"),
        ("appointments", "ownerId, petId, vetId", "Requested/rescheduled time, condition, reason, status"),
        ("products", "review target", "Category, pet type, price, stock, details"),
        ("carearticles", "public content", "Article, video, and FAQ metadata/content"),
        ("adoptionlistings", "shelterId", "Animal profile, images, health/status, care logs"),
        ("adoptioninterests", "listingId, shelterId, adopterId", "Form, reply, and decision lifecycle"),
        ("reviews", "authorId, targetType, targetId", "One-to-five rating and public comment"),
        ("notifications", "userId", "Appointment, vaccine, adoption, product, and system alerts"),
        ("contactmessages", "public submission", "Persisted support/contact inquiries"),
    ], [1800, 3100, 4460])
    callout(doc, "Database choice", "The SRS explicitly permits MongoDB. Consequently, the seed script and Mongoose schemas are the database deliverables; a relational SQL script would not describe the implemented persistence layer accurately.", PALE)

    doc.add_heading("5. Functional requirement traceability", level=1)
    rows = [
        ("FR-01", "Owner account and contact details", "Complete", "Role-aware register/login/profile"),
        ("FR-02", "Multiple pet CRUD and tabbed navigation", "Complete", "Owner pet workspace"),
        ("FR-03", "Gallery, PDFs, X-rays, labs, insurance", "Complete", "Signed Cloudinary upload/view/delete"),
        ("FR-04", "Health CRUD and chronological timeline", "Complete", "Structured records and due dates"),
        ("FR-05", "Products, filters, details, cart quantities", "Complete", "No checkout by SRS exclusion"),
        ("FR-06", "Articles, video, FAQs", "Complete", "Searchable public care library"),
        ("FR-07", "Appointment booking and vet suggestions", "Complete", "Condition/location ranking"),
        ("FR-08", "Vet profile, experience, slots", "Complete", "Registration and profile workspace"),
        ("FR-09", "Booked-pet-only clinical access", "Complete", "Server-side appointment authorization"),
        ("FR-10", "Treatment, labs, prescriptions, follow-up", "Complete", "Structured clinical form"),
        ("FR-11", "Approve, reschedule actual time, complete", "Complete", "Vet schedule controls"),
        ("FR-12", "Shelter listings with images and status", "Complete", "Public catalog + shelter management"),
        ("FR-13", "Feeding, grooming, medical care logs", "Complete", "Per-listing care ledger"),
        ("FR-14", "Adopter form, shelter reply, finalize", "Complete", "Decision lifecycle + alerts"),
        ("FR-15", "Search/sort/filter across discovery", "Complete", "Pets/products/care/vets/adoption"),
        ("FR-16", "Ratings/comments", "Complete", "Owner submission + public feedback"),
        ("FR-17", "About, Contact, map", "Complete", "Dedicated pages + persisted form"),
        ("FR-18", "Role-based access and notifications", "Complete", "Signed sessions + in-app alerts"),
    ]
    matrix(doc, ["ID", "Requirement", "Status", "Evidence"], rows, [850, 3650, 1200, 3660])

    doc.add_heading("6. Security and quality design", level=1)
    bullets(doc, [
        "Passwords are bcrypt-hashed at cost 12; signed session cookies are HTTP-only and SameSite=Lax, with Secure enabled in production.",
        "Database and media credentials remain in ignored environment files and are never returned to the browser.",
        "Upload signatures require authentication; media is restricted to expected formats and 10 MB client limits, and stored URLs must use the configured Cloudinary host.",
        "Veterinarian access to health history requires a matching, non-cancelled appointment for the veterinarian and pet.",
        "Responsive layouts, visible labels, focus indicators, reduced-motion support, semantic headings, native controls, empty states, loading states, and error boundaries support usability and accessibility.",
        "The application uses indexed identifiers, server rendering, reusable domain models, and stateless signed sessions to support scaling and compatibility across current browsers.",
    ])

    doc.add_heading("7. Test data and acceptance verification", level=1)
    matrix(doc, ["Check", "Result", "Evidence"], [
        ("ESLint", "Pass", "npm run lint"),
        ("TypeScript", "Pass", "Production build type-check"),
        ("Next.js production build", "Pass", "26 routes; static and dynamic generation"),
        ("MongoDB seed", "Pass", "Idempotent demo dataset executed successfully"),
        ("Authorization", "Implemented", "Role and ownership gates in pages/actions/routes"),
        ("Secret isolation", "Pass", ".env.local excluded; .env.example contains names only"),
        ("DOCX structural QA", "Pass", "Styles, headings, tables, relationships, and accessibility audited"),
    ], [2600, 1500, 5260])
    add_p(doc, "The seed creates two owner pets, health events, appointments, products, an adoptable animal and care log, adopter interest/reply, reviews, and notifications. Re-running the seed updates known records instead of multiplying them.")

    doc.add_heading("8. Installation and role credentials", level=1)
    steps(doc, [
        "Install Node.js 20 or newer and run npm install in the project directory.",
        "Copy .env.example to .env.local and fill MongoDB, Cloudinary, SESSION_SECRET, and Mailtrap SMTP values.",
        "Run npm run seed to create deterministic demonstration data.",
        "Run npm run dev for local use, or npm run build followed by npm start for a production-mode check.",
    ])
    matrix(doc, ["Role", "Email", "Password"], [
        ("Owner", "owner@furshield.test", "FurShield123!"),
        ("Veterinarian", "vet@furshield.test", "FurShield123!"),
        ("Shelter", "shelter@furshield.test", "FurShield123!"),
    ], [1800, 4260, 3300])
    callout(doc, "Credential hygiene", "These accounts are demonstration-only. Replace the password and rotate all credentials before any public deployment.", "FFF1EC")

    doc.add_heading("9. Release checklist and residual external work", level=1)
    matrix(doc, ["Item", "Local state", "Release action"], [
        ("Application features", "Complete and build-verified", "Run role walkthrough in target environment"),
        ("DOCX visual render", "Blocked: LibreOffice unavailable", "Open once in Word/LibreOffice and visually inspect before submission"),
        ("Email delivery", "Mailtrap SMTP verified", "Replace sandbox SMTP with a production delivery provider before launch"),
        ("Hosted URL", "Not created locally", "Select host, configure secrets, deploy, and smoke-test"),
        ("Mandatory MP4", "Demo script supplied", "Record 7-10 minute walkthrough at 1080p"),
        ("Credential exposure", "User supplied credentials in chat", "Rotate MongoDB and Cloudinary secrets before deployment"),
    ], [2300, 2800, 4260])
    add_p(doc, "This report distinguishes application completion from external operational evidence. A passing local build cannot prove 24/7 availability, a deployed URL, email deliverability, or the existence of a recorded video.")

    path = OUT / "FurShield_Project_Report.docx"
    doc.save(path)
    return path


def build_readme():
    doc = Document()
    configure(doc, "Installation and Demonstration Guide")
    cover(doc, "Operator reference", "FurShield README", "Install, seed, run, verify, and demonstrate every role", [
        ("Audience", "Developers, evaluators, and demonstration operators"),
        ("Runtime", "Node.js 20+; MongoDB Atlas; Cloudinary"),
        ("Package manager", "npm"),
        ("Status", "Production build verified locally"),
    ])
    callout(doc, "Fast path", "Configure the environment, run npm install, run npm run seed, then run npm run dev. Sign in with the three demonstration accounts below.")

    doc.add_heading("1. Prerequisites", level=1)
    bullets(doc, ["Node.js 20 or newer and npm", "MongoDB Atlas connection string", "Cloudinary cloud name, API key, and API secret", "Mailtrap SMTP sandbox credentials for captured email testing"])

    doc.add_heading("2. Environment configuration", level=1)
    add_p(doc, "Copy .env.example to .env.local. Fill the following names; never commit or paste real values into reports or screenshots.")
    matrix(doc, ["Variable", "Required", "Purpose"], [
        ("MONGODB_URI", "Yes", "Application and seed database connection"),
        ("CLOUDINARY_CLOUD_NAME", "Yes", "Media account identifier"),
        ("CLOUDINARY_API_KEY", "Yes", "Signed upload key"),
        ("CLOUDINARY_API_SECRET", "Yes", "Server-only upload signing secret"),
        ("CLOUDINARY_FOLDER", "Yes", "Media namespace, recommended: furshield"),
        ("SESSION_SECRET", "Yes", "Long random session-signing key"),
        ("MAILTRAP_HOST / PORT", "Yes", "SMTP sandbox endpoint"),
        ("MAILTRAP_USER / PASS", "Yes", "Server-only SMTP authentication"),
        ("EMAIL_FROM", "Yes", "Sender identity shown in captured messages"),
    ], [3400, 1200, 4760])

    doc.add_heading("3. Install and run", level=1)
    steps(doc, ["Open a terminal in the project directory.", "Run npm install.", "Run npm run seed.", "Run npm run dev and open http://localhost:3000."])
    add_p(doc, "Production-mode verification: run npm run lint, then npm run build, then npm start.")

    doc.add_heading("4. Demonstration credentials", level=1)
    matrix(doc, ["Role", "Email", "Password", "Primary workspace"], [
        ("Owner", "owner@furshield.test", "FurShield123!", "Pets, records, appointments, adoption, alerts, reviews"),
        ("Veterinarian", "vet@furshield.test", "FurShield123!", "Profile, schedule, booked patients, clinical records"),
        ("Shelter", "shelter@furshield.test", "FurShield123!", "Listings, images, care logs, adopter decisions"),
    ], [1300, 2700, 2000, 3360])

    doc.add_heading("5. Role walkthrough", level=1)
    doc.add_heading("Owner", level=2)
    bullets(doc, ["Switch between multiple pet tabs; search pets; add, edit, and delete a pet.", "Upload an image or PDF, view it, delete it, and maintain vaccine/allergy/illness/treatment/lab/insurance records.", "Search veterinarians by condition/location, request an appointment, and review rescheduled times.", "Filter products and change cart quantities; submit adoption interest; read shelter replies and notifications; publish reviews."])
    doc.add_heading("Veterinarian", level=2)
    bullets(doc, ["Update specialization, experience, and availability slots.", "Approve, reschedule, complete, or cancel appointments.", "Open only booked-pet histories and add symptoms, diagnosis, labs, treatment, prescription, and follow-up notes."])
    doc.add_heading("Shelter", level=2)
    bullets(doc, ["Create adoption profiles, attach images, and update available/pending/adopted status.", "Add feeding, grooming, and medical care logs.", "Review adopter housing/experience forms, reply, and finalize a decision; in-app notification is always created and email is attempted when configured."])
    doc.add_heading("Public", level=2)
    bullets(doc, ["Search and filter care articles/videos/FAQs, vets, adoptable animals, and products.", "Read public ratings/comments, About information, and Contact/map details; submit a contact message."])

    doc.add_heading("6. Verification checklist", level=1)
    matrix(doc, ["Command / check", "Expected result"], [
        ("npm run lint", "No ESLint errors"),
        ("npm run build", "Successful Next.js production build with 26 routes"),
        ("npm run seed", "FurShield demo data is ready"),
        ("Unauthenticated /dashboard", "Redirects to login"),
        ("Vet opens unbooked pet", "Access denied or unavailable"),
        ("Owner/shelter upload", "Signed Cloudinary upload attaches authorized media"),
        ("Shelter decision", "Owner sees persisted reply/status notification"),
    ], [3800, 5560])

    doc.add_heading("7. Troubleshooting", level=1)
    matrix(doc, ["Symptom", "Resolution"], [
        ("Database connection error", "Confirm MONGODB_URI, Atlas network access, and database-user permissions."),
        ("Upload signature or media error", "Confirm all Cloudinary values, folder, file type, and 10 MB size limit."),
        ("Session errors", "Set a long random SESSION_SECRET, clear old cookies, and restart the server."),
        ("No captured email", "In-app alerts still work. Confirm Mailtrap SMTP credentials and inspect the sandbox inbox."),
        ("Build appears slow at page generation", "Allow Atlas-backed dynamic pages time to complete; confirm network connectivity."),
    ], [3000, 6360])

    doc.add_heading("8. Scope and release notes", level=1)
    bullets(doc, [
        "Payments and physical delivery are intentionally not implemented.",
        "Veterinarian credential authentication is intentionally not implemented.",
        "Family account sharing and the chatbot are optional and not enabled.",
        "A public URL requires a deployment target and production secrets.",
        "The mandatory MP4 should follow docs/DEMO_SCRIPT.md and be recorded after deployment or in the local interactive environment.",
        "Rotate the MongoDB and Cloudinary credentials that were shared in chat before public deployment.",
    ])
    path = OUT / "ReadMe.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for item in (build_report(), build_readme()):
        print(item)
